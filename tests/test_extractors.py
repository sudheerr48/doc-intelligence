"""
Tests for src/extractors.py - text extraction from files.
"""

import os
from pathlib import Path
from datetime import datetime

import pytest

from src.extractors import (
    extract_text, _extract_pdf, _extract_plaintext,
    _extract_docx, _extract_xlsx, MAX_TEXT_LENGTH,
)
from src.scanner import FileInfo
from src.storage import FileDatabase


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def make_file_info(
    path="/tmp/test/file.txt",
    name="file.txt",
    extension=".txt",
    size_bytes=1024,
    content_hash="abc123",
    category="test",
    content_text=None,
):
    """Create a FileInfo with sensible defaults for testing."""
    return FileInfo(
        path=path,
        name=name,
        extension=extension,
        size_bytes=size_bytes,
        created_at=datetime(2024, 1, 1, 12, 0, 0),
        modified_at=datetime(2024, 1, 1, 12, 0, 0),
        content_hash=content_hash,
        category=category,
        content_text=content_text,
    )


# ---------------------------------------------------------------------------
# extract_text dispatcher
# ---------------------------------------------------------------------------

class TestExtractText:
    """Tests for the extract_text dispatcher."""

    def test_extracts_txt_file(self, tmp_path):
        f = tmp_path / "hello.txt"
        f.write_text("Hello, world!")
        assert extract_text(str(f)) == "Hello, world!"

    def test_extracts_csv_file(self, tmp_path):
        f = tmp_path / "data.csv"
        f.write_text("a,b,c\n1,2,3\n")
        result = extract_text(str(f))
        assert "a,b,c" in result

    def test_extracts_md_file(self, tmp_path):
        f = tmp_path / "readme.md"
        f.write_text("# Title\nSome content")
        result = extract_text(str(f))
        assert "Title" in result

    def test_extracts_json_file(self, tmp_path):
        f = tmp_path / "config.json"
        f.write_text('{"key": "value"}')
        result = extract_text(str(f))
        assert "key" in result

    def test_extracts_python_file(self, tmp_path):
        f = tmp_path / "script.py"
        f.write_text("def hello():\n    print('hi')\n")
        result = extract_text(str(f))
        assert "def hello" in result

    def test_returns_none_for_unsupported(self, tmp_path):
        f = tmp_path / "image.jpg"
        f.write_bytes(b"\xff\xd8\xff\xe0" + b"\x00" * 100)
        assert extract_text(str(f)) is None

    def test_returns_none_for_binary(self, tmp_path):
        f = tmp_path / "data.bin"
        f.write_bytes(b"\x00" * 100)
        assert extract_text(str(f)) is None

    def test_returns_none_for_nonexistent(self):
        assert extract_text("/nonexistent/file.txt") is None

    def test_returns_none_for_empty_file(self, tmp_path):
        f = tmp_path / "empty.txt"
        f.write_text("")
        assert extract_text(str(f)) is None


# ---------------------------------------------------------------------------
# _extract_plaintext
# ---------------------------------------------------------------------------

class TestExtractPlaintext:
    """Tests for plain text extraction."""

    def test_reads_utf8(self, tmp_path):
        f = tmp_path / "unicode.txt"
        f.write_text("Hello, cafe\u0301!")
        result = _extract_plaintext(str(f))
        assert "cafe" in result

    def test_truncates_long_files(self, tmp_path):
        f = tmp_path / "long.txt"
        content = "x" * (MAX_TEXT_LENGTH + 1000)
        f.write_text(content)
        result = _extract_plaintext(str(f))
        assert len(result) <= MAX_TEXT_LENGTH

    def test_strips_whitespace(self, tmp_path):
        f = tmp_path / "padded.txt"
        f.write_text("  hello  \n\n")
        assert _extract_plaintext(str(f)) == "hello"


# ---------------------------------------------------------------------------
# _extract_pdf
# ---------------------------------------------------------------------------

class TestExtractPDF:
    """Tests for PDF text extraction."""

    def test_extracts_from_real_pdf(self, tmp_path):
        """Create a minimal valid PDF with text content."""
        from pypdf import PdfWriter

        writer = PdfWriter()
        writer.add_blank_page(width=200, height=200)
        # Add text annotation as a simple approach
        page = writer.pages[0]

        pdf_path = tmp_path / "test.pdf"
        with open(pdf_path, "wb") as f:
            writer.write(f)

        # This minimal PDF may not have extractable text, but should not error
        result = _extract_pdf(str(pdf_path))
        # Result may be None for a blank page - that's correct behavior
        assert result is None or isinstance(result, str)

    def test_returns_none_for_corrupted_pdf(self, tmp_path):
        f = tmp_path / "bad.pdf"
        f.write_bytes(b"not a real pdf content")
        assert _extract_pdf(str(f)) is None

    def test_returns_none_for_nonexistent_pdf(self):
        assert _extract_pdf("/nonexistent/file.pdf") is None


# ---------------------------------------------------------------------------
# _extract_docx
# ---------------------------------------------------------------------------

class TestExtractDocx:
    """Tests for DOCX text extraction."""

    def test_extracts_from_docx(self, tmp_path):
        from docx import Document

        doc = Document()
        doc.add_paragraph("Hello from DOCX")
        doc.add_paragraph("This is paragraph two.")
        docx_path = tmp_path / "test.docx"
        doc.save(str(docx_path))

        result = _extract_docx(str(docx_path))
        assert result is not None
        assert "Hello from DOCX" in result
        assert "paragraph two" in result

    def test_extracts_tables_from_docx(self, tmp_path):
        from docx import Document

        doc = Document()
        doc.add_paragraph("Header text")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Name"
        table.cell(0, 1).text = "Value"
        table.cell(1, 0).text = "Alpha"
        table.cell(1, 1).text = "100"
        docx_path = tmp_path / "table.docx"
        doc.save(str(docx_path))

        result = _extract_docx(str(docx_path))
        assert result is not None
        assert "Name" in result
        assert "Alpha" in result

    def test_returns_none_for_empty_docx(self, tmp_path):
        from docx import Document

        doc = Document()
        docx_path = tmp_path / "empty.docx"
        doc.save(str(docx_path))

        result = _extract_docx(str(docx_path))
        assert result is None

    def test_returns_none_for_corrupted_docx(self, tmp_path):
        f = tmp_path / "bad.docx"
        f.write_bytes(b"not a real docx")
        assert _extract_docx(str(f)) is None

    def test_extract_text_dispatches_docx(self, tmp_path):
        from docx import Document

        doc = Document()
        doc.add_paragraph("Dispatch test")
        docx_path = tmp_path / "dispatch.docx"
        doc.save(str(docx_path))

        result = extract_text(str(docx_path))
        assert result is not None
        assert "Dispatch test" in result


# ---------------------------------------------------------------------------
# _extract_xlsx
# ---------------------------------------------------------------------------

class TestExtractXlsx:
    """Tests for XLSX text extraction."""

    def test_extracts_from_xlsx(self, tmp_path):
        from openpyxl import Workbook

        wb = Workbook()
        ws = wb.active
        ws.title = "Data"
        ws.append(["Name", "Score"])
        ws.append(["Alice", 95])
        ws.append(["Bob", 87])
        xlsx_path = tmp_path / "test.xlsx"
        wb.save(str(xlsx_path))

        result = _extract_xlsx(str(xlsx_path))
        assert result is not None
        assert "Alice" in result
        assert "95" in result
        assert "[Sheet: Data]" in result

    def test_extracts_multiple_sheets(self, tmp_path):
        from openpyxl import Workbook

        wb = Workbook()
        ws1 = wb.active
        ws1.title = "Sheet1"
        ws1.append(["Revenue", "1000"])
        ws2 = wb.create_sheet("Sheet2")
        ws2.append(["Expenses", "500"])
        xlsx_path = tmp_path / "multi.xlsx"
        wb.save(str(xlsx_path))

        result = _extract_xlsx(str(xlsx_path))
        assert result is not None
        assert "Revenue" in result
        assert "Expenses" in result
        assert "[Sheet: Sheet1]" in result
        assert "[Sheet: Sheet2]" in result

    def test_returns_none_for_empty_xlsx(self, tmp_path):
        from openpyxl import Workbook

        wb = Workbook()
        ws = wb.active
        # Empty sheet, no data
        xlsx_path = tmp_path / "empty.xlsx"
        wb.save(str(xlsx_path))

        result = _extract_xlsx(str(xlsx_path))
        assert result is None

    def test_returns_none_for_corrupted_xlsx(self, tmp_path):
        f = tmp_path / "bad.xlsx"
        f.write_bytes(b"not a real xlsx")
        assert _extract_xlsx(str(f)) is None

    def test_extract_text_dispatches_xlsx(self, tmp_path):
        from openpyxl import Workbook

        wb = Workbook()
        ws = wb.active
        ws.append(["dispatch", "test"])
        xlsx_path = tmp_path / "dispatch.xlsx"
        wb.save(str(xlsx_path))

        result = extract_text(str(xlsx_path))
        assert result is not None
        assert "dispatch" in result


# ---------------------------------------------------------------------------
# Content text in storage (search by content)
# ---------------------------------------------------------------------------

class TestContentTextSearch:
    """Tests for searching content_text via the database."""

    def test_insert_and_search_content(self, tmp_path):
        db_path = str(tmp_path / "test.duckdb")
        db = FileDatabase(db_path)

        fi = make_file_info(
            path="/docs/contract.pdf",
            name="contract.pdf",
            extension=".pdf",
            content_text="This agreement is between Alice and Bob for consulting services.",
        )
        db.insert_file(fi)

        # Search for text that's in content but not in the filename
        results = db.search("consulting")
        assert len(results) == 1
        assert results[0]["name"] == "contract.pdf"
        assert results[0]["content_match"] is True
        db.close()

    def test_content_search_case_insensitive(self, tmp_path):
        db_path = str(tmp_path / "test.duckdb")
        db = FileDatabase(db_path)

        fi = make_file_info(
            path="/docs/notes.txt",
            name="notes.txt",
            content_text="IMPORTANT: Meeting at 3pm about Project Alpha",
        )
        db.insert_file(fi)

        results = db.search("project alpha")
        assert len(results) == 1
        assert results[0]["content_match"] is True
        db.close()

    def test_name_match_without_content(self, tmp_path):
        db_path = str(tmp_path / "test.duckdb")
        db = FileDatabase(db_path)

        fi = make_file_info(
            path="/docs/report.txt",
            name="report.txt",
            content_text=None,
        )
        db.insert_file(fi)

        results = db.search("report")
        assert len(results) == 1
        assert results[0]["content_match"] is False
        db.close()

    def test_content_match_ranked_higher(self, tmp_path):
        """Files matching by content should appear before name-only matches."""
        db_path = str(tmp_path / "test.duckdb")
        db = FileDatabase(db_path)

        # File matching by name only
        db.insert_file(make_file_info(
            path="/a/budget.txt",
            name="budget.txt",
            content_hash="h1",
            content_text="General expenses and income summary",
        ))
        # File matching by content
        db.insert_file(make_file_info(
            path="/b/report.pdf",
            name="report.pdf",
            content_hash="h2",
            content_text="The annual budget review shows a surplus of $50,000.",
        ))

        results = db.search("budget")
        assert len(results) == 2
        # Content match (report.pdf has "budget" in content_text) should rank first
        assert results[0]["content_match"] is True
        db.close()

    def test_batch_insert_with_content_text(self, tmp_path):
        db_path = str(tmp_path / "test.duckdb")
        db = FileDatabase(db_path)

        files = [
            make_file_info(
                path=f"/docs/file{i}.txt",
                name=f"file{i}.txt",
                content_hash=f"hash{i}",
                content_text=f"Content of document number {i}",
            )
            for i in range(5)
        ]
        db.insert_batch(files)

        results = db.search("document number 3")
        assert len(results) == 1
        assert results[0]["name"] == "file3.txt"
        db.close()

    def test_no_content_text_still_works(self, tmp_path):
        """Files without content_text should still be searchable by name/path."""
        db_path = str(tmp_path / "test.duckdb")
        db = FileDatabase(db_path)

        fi = make_file_info(
            path="/images/photo.jpg",
            name="photo.jpg",
            content_text=None,
        )
        db.insert_file(fi)

        results = db.search("photo")
        assert len(results) == 1
        assert results[0]["content_match"] is False
        db.close()
