"""
Text Extraction Module
Extracts text content from files for full-text search.
Supports PDF, DOCX, XLSX, Markdown, and plain text files.
"""

from pathlib import Path
from typing import Optional

# Maximum text to store per file (64KB)
MAX_TEXT_LENGTH = 65536

# Plain text extensions handled by _extract_plaintext
_PLAINTEXT_EXTS = frozenset((
    ".txt", ".csv", ".md", ".json", ".xml", ".html", ".htm",
    ".log", ".yaml", ".yml", ".toml", ".ini", ".cfg", ".conf",
    ".py", ".js", ".ts", ".java", ".c", ".cpp", ".h", ".go",
    ".rs", ".rb", ".sh", ".bat", ".sql", ".rst", ".tex",
    ".r", ".m", ".swift", ".kt", ".scala", ".pl", ".lua",
    ".ps1", ".dockerfile", ".makefile",
))


def extract_text(file_path: str) -> Optional[str]:
    """
    Extract text content from a file based on its extension.

    Args:
        file_path: Path to the file

    Returns:
        Extracted text, or None if extraction fails or is unsupported
    """
    ext = Path(file_path).suffix.lower()
    name = Path(file_path).name.lower()

    try:
        if ext == ".pdf":
            return _extract_pdf(file_path)
        elif ext == ".docx":
            return _extract_docx(file_path)
        elif ext == ".xlsx":
            return _extract_xlsx(file_path)
        elif ext in _PLAINTEXT_EXTS or name in ("makefile", "dockerfile", "rakefile", "gemfile"):
            return _extract_plaintext(file_path)
        else:
            return None
    except Exception:
        return None


def _extract_pdf(file_path: str) -> Optional[str]:
    """Extract text from a PDF file using pypdf."""
    try:
        from pypdf import PdfReader

        reader = PdfReader(file_path)
        pages = []
        total_len = 0

        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
                total_len += len(text)
                if total_len >= MAX_TEXT_LENGTH:
                    break

        if not pages:
            return None

        full_text = "\n".join(pages)
        if len(full_text) > MAX_TEXT_LENGTH:
            full_text = full_text[:MAX_TEXT_LENGTH]

        return full_text.strip() or None
    except BaseException:
        # BaseException catches import failures (e.g. missing cffi)
        # and Rust panics from cryptography bindings
        return None


def _extract_docx(file_path: str) -> Optional[str]:
    """Extract text from a DOCX file using python-docx."""
    try:
        from docx import Document

        doc = Document(file_path)
        paragraphs = []
        total_len = 0

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
                total_len += len(text) + 1
                if total_len >= MAX_TEXT_LENGTH:
                    break

        # Also extract text from tables
        if total_len < MAX_TEXT_LENGTH:
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        line = " | ".join(cells)
                        paragraphs.append(line)
                        total_len += len(line) + 1
                        if total_len >= MAX_TEXT_LENGTH:
                            break
                if total_len >= MAX_TEXT_LENGTH:
                    break

        if not paragraphs:
            return None

        full_text = "\n".join(paragraphs)
        if len(full_text) > MAX_TEXT_LENGTH:
            full_text = full_text[:MAX_TEXT_LENGTH]

        return full_text.strip() or None
    except Exception:
        return None


def _extract_xlsx(file_path: str) -> Optional[str]:
    """Extract text from an XLSX file using openpyxl."""
    try:
        from openpyxl import load_workbook

        wb = load_workbook(file_path, read_only=True, data_only=True)
        lines = []
        total_len = 0

        for sheet in wb.sheetnames:
            ws = wb[sheet]
            lines.append(f"[Sheet: {sheet}]")

            for row in ws.iter_rows(values_only=True):
                cells = [str(c).strip() for c in row if c is not None]
                if cells:
                    line = " | ".join(cells)
                    lines.append(line)
                    total_len += len(line) + 1
                    if total_len >= MAX_TEXT_LENGTH:
                        break
            if total_len >= MAX_TEXT_LENGTH:
                break

        wb.close()

        if len(lines) <= 1:  # Only sheet header, no data
            return None

        full_text = "\n".join(lines)
        if len(full_text) > MAX_TEXT_LENGTH:
            full_text = full_text[:MAX_TEXT_LENGTH]

        return full_text.strip() or None
    except Exception:
        return None


def _extract_plaintext(file_path: str) -> Optional[str]:
    """Extract text from a plain text file."""
    try:
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            text = f.read(MAX_TEXT_LENGTH)
        return text.strip() or None
    except Exception:
        return None
