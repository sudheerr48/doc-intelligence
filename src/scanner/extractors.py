"""
Text extraction from files — PDF, DOCX, XLSX, and plain text.
"""

from pathlib import Path
from typing import Optional

MAX_TEXT_LENGTH = 65536

_PLAINTEXT_EXTS = frozenset((
    ".txt", ".csv", ".md", ".json", ".xml", ".html", ".htm",
    ".log", ".yaml", ".yml", ".toml", ".ini", ".cfg", ".conf",
    ".py", ".js", ".ts", ".java", ".c", ".cpp", ".h", ".go",
    ".rs", ".rb", ".sh", ".bat", ".sql", ".rst", ".tex",
    ".r", ".m", ".swift", ".kt", ".scala", ".pl", ".lua",
    ".ps1", ".dockerfile", ".makefile",
))


def extract_text(file_path: str) -> Optional[str]:
    """Extract text content from a file based on its extension."""
    ext = Path(file_path).suffix.lower()
    name = Path(file_path).name.lower()

    try:
        if ext == ".pdf":
            return _extract_pdf(file_path)
        elif ext == ".docx":
            return _extract_docx(file_path)
        elif ext == ".xlsx":
            return _extract_xlsx(file_path)
        elif ext in _PLAINTEXT_EXTS or name in (
            "makefile", "dockerfile", "rakefile", "gemfile",
        ):
            return _extract_plaintext(file_path)
        else:
            return None
    except Exception:
        return None


def _extract_pdf(file_path: str) -> Optional[str]:
    try:
        from pypdf import PdfReader
        reader = PdfReader(file_path)
        pages = []
        total_len = 0
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
                total_len += len(text)
                if total_len >= MAX_TEXT_LENGTH:
                    break
        if not pages:
            return None
        full_text = "\n".join(pages)
        if len(full_text) > MAX_TEXT_LENGTH:
            full_text = full_text[:MAX_TEXT_LENGTH]
        return full_text.strip() or None
    except BaseException:
        return None


def _extract_docx(file_path: str) -> Optional[str]:
    try:
        from docx import Document
        doc = Document(file_path)
        paragraphs = []
        total_len = 0
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
                total_len += len(text) + 1
                if total_len >= MAX_TEXT_LENGTH:
                    break
        if total_len < MAX_TEXT_LENGTH:
            for table in doc.tables:
                for row in table.rows:
                    cells = [
                        cell.text.strip()
                        for cell in row.cells if cell.text.strip()
                    ]
                    if cells:
                        line = " | ".join(cells)
                        paragraphs.append(line)
                        total_len += len(line) + 1
                        if total_len >= MAX_TEXT_LENGTH:
                            break
                if total_len >= MAX_TEXT_LENGTH:
                    break
        if not paragraphs:
            return None
        full_text = "\n".join(paragraphs)
        if len(full_text) > MAX_TEXT_LENGTH:
            full_text = full_text[:MAX_TEXT_LENGTH]
        return full_text.strip() or None
    except Exception:
        return None


def _extract_xlsx(file_path: str) -> Optional[str]:
    try:
        from openpyxl import load_workbook
        wb = load_workbook(file_path, read_only=True, data_only=True)
        lines = []
        total_len = 0
        for sheet in wb.sheetnames:
            ws = wb[sheet]
            lines.append(f"[Sheet: {sheet}]")
            for row in ws.iter_rows(values_only=True):
                cells = [str(c).strip() for c in row if c is not None]
                if cells:
                    line = " | ".join(cells)
                    lines.append(line)
                    total_len += len(line) + 1
                    if total_len >= MAX_TEXT_LENGTH:
                        break
            if total_len >= MAX_TEXT_LENGTH:
                break
        wb.close()
        if len(lines) <= 1:
            return None
        full_text = "\n".join(lines)
        if len(full_text) > MAX_TEXT_LENGTH:
            full_text = full_text[:MAX_TEXT_LENGTH]
        return full_text.strip() or None
    except Exception:
        return None


def _extract_plaintext(file_path: str) -> Optional[str]:
    try:
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            text = f.read(MAX_TEXT_LENGTH)
        return text.strip() or None
    except Exception:
        return None
